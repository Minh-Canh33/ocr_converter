"""
==============================================================
  OCR App - Đọc văn bản từ ảnh & Xuất DOCX/PDF
==============================================================

Cài thư viện trước khi chạy:
  pip install pytesseract pillow python-docx reportlab opencv-python

Cài Tesseract OCR engine:
  - Windows : https://github.com/UB-Mannheim/tesseract/wiki
              Thêm vào PATH hoặc khai báo: pytesseract.pytesseract.tesseract_cmd = r"C:\...\tesseract.exe"
  - Ubuntu  : sudo apt install tesseract-ocr tesseract-ocr-vie
  - macOS   : brew install tesseract tesseract-lang

Chạy:
  python ocr_app.py
==============================================================
"""

import os
import sys
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext, ttk


# ── Kiểm tra thư viện bắt buộc ────────────────────────────────────────────────
MISSING = []
try:
    import pytesseract
except ImportError:
    MISSING.append("pytesseract")

try:
    from PIL import Image, ImageTk, ImageEnhance, ImageFilter
except ImportError:
    MISSING.append("pillow")

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    MISSING.append("python-docx")

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
except ImportError:
    MISSING.append("reportlab")

try:
    import cv2
    import numpy as np
    HAS_CV2 = True
except ImportError:
    HAS_CV2 = False  # Bonus – không bắt buộc

if MISSING:
    print(f"[LỖI] Thiếu thư viện: {', '.join(MISSING)}")
    print("Chạy: pip install " + " ".join(MISSING))
    sys.exit(1)


# ══════════════════════════════════════════════════════════════════════════════
#  1. ĐỌC ẢNH
# ══════════════════════════════════════════════════════════════════════════════

def load_image(path: str) -> Image.Image:
    """Đọc file ảnh (jpg/png/bmp/tiff…) và trả về PIL Image."""
    supported = (".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".tif", ".webp")
    ext = os.path.splitext(path)[1].lower()
    if ext not in supported:
        raise ValueError(f"Định dạng không được hỗ trợ: {ext}\nHỗ trợ: {', '.join(supported)}")
    if not os.path.exists(path):
        raise FileNotFoundError(f"Không tìm thấy file: {path}")
    return Image.open(path).convert("RGB")


# ══════════════════════════════════════════════════════════════════════════════
#  2. TIỀN XỬ LÝ ẢNH (Bonus – OpenCV)
# ══════════════════════════════════════════════════════════════════════════════

def preprocess_image(pil_img: Image.Image, use_cv2: bool = True) -> Image.Image:
    """
    Tăng độ chính xác OCR bằng cách:
      - Chuyển grayscale
      - Tăng độ tương phản
      - Làm sắc nét
      - Ngưỡng hóa (thresholding) nếu có OpenCV
    """
    if use_cv2 and HAS_CV2:
        # Chuyển PIL → numpy
        img_np = np.array(pil_img)
        gray = cv2.cvtColor(img_np, cv2.COLOR_RGB2GRAY)

        # Loại bỏ nhiễu nhẹ
        denoised = cv2.fastNlMeansDenoising(gray, h=10)

        # Adaptive threshold – tốt cho ảnh chụp tài liệu
        thresh = cv2.adaptiveThreshold(
            denoised, 255,
            cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
            cv2.THRESH_BINARY, 31, 10
        )

        # Phóng to nếu ảnh nhỏ (< 1000px chiều rộng)
        h, w = thresh.shape
        if w < 1000:
            scale = 1000 / w
            thresh = cv2.resize(thresh, None, fx=scale, fy=scale,
                                interpolation=cv2.INTER_CUBIC)

        return Image.fromarray(thresh)

    else:
        # Fallback thuần PIL
        img = pil_img.convert("L")                          # grayscale
        img = ImageEnhance.Contrast(img).enhance(2.0)       # tăng tương phản
        img = img.filter(ImageFilter.SHARPEN)               # làm sắc nét
        return img


# ══════════════════════════════════════════════════════════════════════════════
#  3. OCR
# ══════════════════════════════════════════════════════════════════════════════

def run_ocr(pil_img: Image.Image, lang: str = "eng") -> str:
    """
    Chạy Tesseract OCR.
    lang: 'eng' (Anh) | 'vie' (Việt) | 'eng+vie' (cả hai)
    """
    config = "--oem 3 --psm 6"   # LSTM engine + assume uniform text block
    try:
        text = pytesseract.image_to_string(pil_img, lang=lang, config=config)
        return text.strip()
    except pytesseract.TesseractNotFoundError:
        raise RuntimeError(
            "Không tìm thấy Tesseract.\n"
            "Cài đặt tại: https://github.com/UB-Mannheim/tesseract/wiki\n"
            "Hoặc: sudo apt install tesseract-ocr"
        )
    except Exception as e:
        raise RuntimeError(f"OCR thất bại: {e}")


# ══════════════════════════════════════════════════════════════════════════════
#  4. LƯU DOCX
# ══════════════════════════════════════════════════════════════════════════════

def save_docx(text: str, output_path: str, image_path: str = "") -> str:
    """Lưu văn bản OCR ra file .docx với định dạng cơ bản."""
    doc = DocxDocument()

    # Tiêu đề
    title = doc.add_heading("Kết quả OCR", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    # Thông tin nguồn
    if image_path:
        info = doc.add_paragraph()
        info.add_run("📁 Nguồn: ").bold = True
        info.add_run(os.path.basename(image_path))
        info.add_run(f"\n📅 Thư mục: ").bold = True
        info.add_run(os.path.dirname(os.path.abspath(image_path)))

    doc.add_paragraph()  # khoảng cách

    # Nội dung văn bản
    heading = doc.add_heading("Nội dung văn bản", level=2)
    heading.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    for line in text.split("\n"):
        p = doc.add_paragraph(line)
        p.style.font.size = Pt(12)

    # Footer note
    doc.add_paragraph()
    note = doc.add_paragraph()
    note.add_run("✨ Tạo bởi OCR App – pytesseract").italic = True
    note.runs[0].font.size = Pt(9)
    note.runs[0].font.color.rgb = RGBColor(0x70, 0x70, 0x70)

    doc.save(output_path)
    return output_path


# ══════════════════════════════════════════════════════════════════════════════
#  5. LƯU PDF
# ══════════════════════════════════════════════════════════════════════════════

def save_pdf(text: str, output_path: str, image_path: str = "") -> str:
    """Lưu văn bản OCR ra file .pdf bằng ReportLab."""
    doc = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=25 * mm,
        bottomMargin=20 * mm,
    )

    styles = getSampleStyleSheet()

    # Thử đăng ký font hỗ trợ Unicode (tiếng Việt)
    _register_unicode_font(styles)

    title_style = ParagraphStyle(
        "ocr_title",
        parent=styles["Title"],
        fontSize=18,
        textColor="#1F497D",
        spaceAfter=6,
        fontName=styles["Title"].fontName,
    )
    body_style = ParagraphStyle(
        "ocr_body",
        parent=styles["Normal"],
        fontSize=11,
        leading=16,
        spaceAfter=2,
        fontName=styles["Normal"].fontName,
    )
    meta_style = ParagraphStyle(
        "ocr_meta",
        parent=styles["Italic"],
        fontSize=9,
        textColor="#707070",
        fontName=styles["Italic"].fontName,
    )

    story = []
    story.append(Paragraph("Kết quả OCR", title_style))
    story.append(Spacer(1, 4 * mm))

    if image_path:
        story.append(Paragraph(
            f"<b>Nguồn:</b> {os.path.basename(image_path)}", body_style
        ))
        story.append(Spacer(1, 4 * mm))

    story.append(Paragraph("<b>Nội dung văn bản</b>", body_style))
    story.append(Spacer(1, 2 * mm))

    for line in text.split("\n"):
        safe = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        if safe.strip():
            story.append(Paragraph(safe, body_style))
        else:
            story.append(Spacer(1, 3 * mm))

    story.append(Spacer(1, 6 * mm))
    story.append(Paragraph("✨ Tạo bởi OCR App – pytesseract &amp; ReportLab", meta_style))

    doc.build(story)
    return output_path


def _register_unicode_font(styles):
    """Thử đăng ký DejaVuSans để hiển thị Unicode / tiếng Việt trong PDF."""
    candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",         # Linux
        "/Library/Fonts/Arial Unicode MS.ttf",                      # macOS
        "C:/Windows/Fonts/arial.ttf",                               # Windows
        "C:/Windows/Fonts/calibri.ttf",
    ]
    for path in candidates:
        if os.path.exists(path):
            try:
                font_name = "UniFont"
                pdfmetrics.registerFont(TTFont(font_name, path))
                styles["Normal"].fontName = font_name
                styles["Title"].fontName = font_name
                styles["Italic"].fontName = font_name
                return
            except Exception:
                pass


# ══════════════════════════════════════════════════════════════════════════════
#  6. GIAO DIỆN TKINTER
# ══════════════════════════════════════════════════════════════════════════════

class OCRApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("📄 OCR App – Đọc văn bản từ ảnh")
        self.geometry("950x700")
        self.minsize(750, 550)
        self.configure(bg="#F0F4F8")

        self._image_path = ""
        self._pil_image = None
        self._ocr_text = ""

        self._build_ui()

    # ── Xây dựng UI ───────────────────────────────────────────────────────────

    def _build_ui(self):
        # ── Top toolbar ───────────────────────────────────────────────────────
        toolbar = tk.Frame(self, bg="#1F497D", pady=6)
        toolbar.pack(fill="x")

        tk.Label(toolbar, text="📷 OCR App", font=("Segoe UI", 15, "bold"),
                 bg="#1F497D", fg="white").pack(side="left", padx=14)

        # Nút chọn ảnh
        self._btn_open = tk.Button(
            toolbar, text="📂  Chọn ảnh",
            font=("Segoe UI", 10, "bold"),
            bg="#2E74B5", fg="white", activebackground="#1F497D",
            relief="flat", padx=12, pady=4, cursor="hand2",
            command=self._open_image
        )
        self._btn_open.pack(side="left", padx=8)

        # Ngôn ngữ OCR
        tk.Label(toolbar, text="Ngôn ngữ:", bg="#1F497D", fg="#CCDDF0",
                 font=("Segoe UI", 10)).pack(side="left")
        self._lang_var = tk.StringVar(value="eng")
        lang_cb = ttk.Combobox(
            toolbar, textvariable=self._lang_var,
            values=["eng", "vie", "eng+vie"],
            width=9, state="readonly", font=("Segoe UI", 10)
        )
        lang_cb.pack(side="left", padx=(2, 14))

        # Tiền xử lý
        self._preprocess_var = tk.BooleanVar(value=True)
        tk.Checkbutton(
            toolbar, text="Tiền xử lý ảnh",
            variable=self._preprocess_var,
            bg="#1F497D", fg="white", selectcolor="#2E74B5",
            activebackground="#1F497D", activeforeground="white",
            font=("Segoe UI", 10)
        ).pack(side="left", padx=4)

        # Nút OCR
        self._btn_ocr = tk.Button(
            toolbar, text="🔍  Nhận diện",
            font=("Segoe UI", 10, "bold"),
            bg="#217346", fg="white", activebackground="#1A5C38",
            relief="flat", padx=12, pady=4, cursor="hand2",
            state="disabled", command=self._do_ocr
        )
        self._btn_ocr.pack(side="left", padx=4)

        # Nút xuất
        self._btn_docx = tk.Button(
            toolbar, text="💾  Lưu DOCX",
            font=("Segoe UI", 10, "bold"),
            bg="#7B3F9E", fg="white", activebackground="#5A2D77",
            relief="flat", padx=12, pady=4, cursor="hand2",
            state="disabled", command=self._save_docx
        )
        self._btn_docx.pack(side="right", padx=4)

        self._btn_pdf = tk.Button(
            toolbar, text="📑  Lưu PDF",
            font=("Segoe UI", 10, "bold"),
            bg="#C84B11", fg="white", activebackground="#9C3A0D",
            relief="flat", padx=12, pady=4, cursor="hand2",
            state="disabled", command=self._save_pdf
        )
        self._btn_pdf.pack(side="right", padx=4)

        # ── Content pane ──────────────────────────────────────────────────────
        pane = tk.PanedWindow(self, orient="horizontal", bg="#D6E4F0",
                              sashwidth=6, sashrelief="groove")
        pane.pack(fill="both", expand=True, padx=6, pady=6)

        # Left – preview ảnh
        left = tk.Frame(pane, bg="#F0F4F8")
        pane.add(left, minsize=250)

        tk.Label(left, text="Xem trước ảnh", bg="#2E74B5", fg="white",
                 font=("Segoe UI", 10, "bold"), pady=4).pack(fill="x")

        self._img_label = tk.Label(
            left, text="Chưa chọn ảnh\n\nNhấn '📂 Chọn ảnh'",
            bg="#DAEAF7", fg="#555", font=("Segoe UI", 10),
            relief="sunken", anchor="center"
        )
        self._img_label.pack(fill="both", expand=True, padx=4, pady=4)

        # Right – text result
        right = tk.Frame(pane, bg="#F0F4F8")
        pane.add(right, minsize=350)

        tk.Label(right, text="Kết quả văn bản OCR", bg="#217346", fg="white",
                 font=("Segoe UI", 10, "bold"), pady=4).pack(fill="x")

        self._text_box = scrolledtext.ScrolledText(
            right, font=("Consolas", 11), wrap="word",
            bg="#FFFEF5", fg="#222", insertbackground="#1F497D",
            relief="sunken", bd=2
        )
        self._text_box.pack(fill="both", expand=True, padx=4, pady=4)

        # ── Status bar ────────────────────────────────────────────────────────
        self._status = tk.StringVar(value="Sẵn sàng. Chọn một file ảnh để bắt đầu.")
        status_bar = tk.Label(
            self, textvariable=self._status,
            bg="#1F497D", fg="#CCDDF0",
            font=("Segoe UI", 9), anchor="w", pady=4, padx=10
        )
        status_bar.pack(fill="x", side="bottom")

    # ── Xử lý sự kiện ────────────────────────────────────────────────────────

    def _open_image(self):
        path = filedialog.askopenfilename(
            title="Chọn file ảnh",
            filetypes=[
                ("Ảnh", "*.jpg *.jpeg *.png *.bmp *.tiff *.tif *.webp"),
                ("Tất cả", "*.*"),
            ]
        )
        if not path:
            return

        try:
            self._pil_image = load_image(path)
            self._image_path = path
            self._show_preview(self._pil_image)
            self._btn_ocr.config(state="normal")
            self._status.set(f"Đã tải: {os.path.basename(path)}  ({self._pil_image.width}×{self._pil_image.height}px)")
            self._text_box.delete("1.0", "end")
            self._btn_docx.config(state="disabled")
            self._btn_pdf.config(state="disabled")
        except Exception as e:
            messagebox.showerror("Lỗi tải ảnh", str(e))

    def _show_preview(self, pil_img: Image.Image):
        """Thu nhỏ ảnh vừa với khung xem trước."""
        self._img_label.update_idletasks()
        max_w = max(self._img_label.winfo_width() - 8, 200)
        max_h = max(self._img_label.winfo_height() - 8, 200)

        img_copy = pil_img.copy()
        img_copy.thumbnail((max_w, max_h), Image.LANCZOS)
        photo = ImageTk.PhotoImage(img_copy)
        self._img_label.config(image=photo, text="")
        self._img_label.image = photo  # giữ reference

    def _do_ocr(self):
        if self._pil_image is None:
            return

        self._status.set("⏳ Đang nhận diện văn bản…")
        self.update()

        try:
            img = self._pil_image
            if self._preprocess_var.get():
                self._status.set("⚙️ Đang tiền xử lý ảnh…")
                self.update()
                img = preprocess_image(img, use_cv2=HAS_CV2)

            lang = self._lang_var.get()
            self._status.set(f"🔍 Đang OCR (ngôn ngữ: {lang})…")
            self.update()

            text = run_ocr(img, lang=lang)
            self._ocr_text = text

            self._text_box.delete("1.0", "end")
            if text:
                self._text_box.insert("1.0", text)
                self._btn_docx.config(state="normal")
                self._btn_pdf.config(state="normal")
                word_count = len(text.split())
                self._status.set(f"✅ Hoàn thành! Nhận diện {word_count} từ / {len(text)} ký tự.")
            else:
                self._text_box.insert("1.0", "(Không tìm thấy văn bản trong ảnh)")
                self._status.set("⚠️ Không tìm thấy văn bản. Thử thay đổi ngôn ngữ hoặc tắt tiền xử lý.")

        except RuntimeError as e:
            messagebox.showerror("Lỗi OCR", str(e))
            self._status.set("❌ OCR thất bại.")

    def _save_docx(self):
        if not self._ocr_text:
            messagebox.showwarning("Chưa có nội dung", "Vui lòng nhận diện văn bản trước.")
            return

        default_name = os.path.splitext(os.path.basename(self._image_path))[0] + "_ocr.docx"
        path = filedialog.asksaveasfilename(
            title="Lưu file DOCX",
            defaultextension=".docx",
            initialfile=default_name,
            filetypes=[("Word Document", "*.docx")]
        )
        if not path:
            return

        try:
            save_docx(self._ocr_text, path, self._image_path)
            self._status.set(f"✅ Đã lưu DOCX: {path}")
            messagebox.showinfo("Lưu thành công", f"File đã được lưu:\n{path}")
        except Exception as e:
            messagebox.showerror("Lỗi lưu DOCX", str(e))

    def _save_pdf(self):
        if not self._ocr_text:
            messagebox.showwarning("Chưa có nội dung", "Vui lòng nhận diện văn bản trước.")
            return

        default_name = os.path.splitext(os.path.basename(self._image_path))[0] + "_ocr.pdf"
        path = filedialog.asksaveasfilename(
            title="Lưu file PDF",
            defaultextension=".pdf",
            initialfile=default_name,
            filetypes=[("PDF", "*.pdf")]
        )
        if not path:
            return

        try:
            save_pdf(self._ocr_text, path, self._image_path)
            self._status.set(f"✅ Đã lưu PDF: {path}")
            messagebox.showinfo("Lưu thành công", f"File đã được lưu:\n{path}")
        except Exception as e:
            messagebox.showerror("Lỗi lưu PDF", str(e))


# ══════════════════════════════════════════════════════════════════════════════
#  7. DEMO CLI (chạy không cần giao diện)
# ══════════════════════════════════════════════════════════════════════════════

def demo_cli(image_path: str, lang: str = "eng"):
    """Ví dụ chạy thử từ command line."""
    print(f"\n{'='*55}")
    print("  OCR App – Demo CLI")
    print(f"{'='*55}")

    print(f"\n[1] Đọc ảnh: {image_path}")
    img = load_image(image_path)
    print(f"    Kích thước: {img.width}×{img.height}px")

    print("\n[2] Tiền xử lý ảnh…")
    processed = preprocess_image(img, use_cv2=HAS_CV2)
    cv2_note = "OpenCV" if HAS_CV2 else "PIL (không có OpenCV)"
    print(f"    Dùng: {cv2_note}")

    print(f"\n[3] Chạy OCR (ngôn ngữ: {lang})…")
    text = run_ocr(processed, lang=lang)
    print("\n── Kết quả văn bản ──")
    print(text or "(Không tìm thấy văn bản)")

    base = os.path.splitext(image_path)[0]

    print("\n[4] Lưu DOCX…")
    docx_path = save_docx(text, base + "_ocr.docx", image_path)
    print(f"    → {docx_path}")

    print("\n[5] Lưu PDF…")
    pdf_path = save_pdf(text, base + "_ocr.pdf", image_path)
    print(f"    → {pdf_path}")

    print("\n✅ Hoàn thành!\n")


# ══════════════════════════════════════════════════════════════════════════════
#  MAIN
# ══════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    if len(sys.argv) == 2:
        # Chế độ CLI: python ocr_app.py path/to/image.jpg
        demo_cli(sys.argv[1], lang="eng+vie")
    elif len(sys.argv) == 3:
        # python ocr_app.py path/to/image.jpg vie
        demo_cli(sys.argv[1], lang=sys.argv[2])
    else:
        # Chế độ GUI
        app = OCRApp()
        app.mainloop()
